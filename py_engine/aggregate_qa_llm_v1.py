#!/usr/bin/env python3
"""
aggregate_qa_llm_v1.py

Quick-test grounded Q&A over {Manager_ID}_aggregate.json.

Behavior:
- Uses ONLY the aggregate JSON as evidence.
- User selects Manager_ID + provides User_ID.
- Loads {INPUT_DIR}\\{Manager_ID}_aggregate.json.
- Interactive query loop:
    - Retrieve top-K relevant bullets (TF-IDF; deterministic re-ranking with recency + source priority).
    - Ask the LLM to decide sufficiency and write a professional answer grounded ONLY in retrieved bullets.
    - If insufficient OR user thumbs-down: append record to:
         {OUTPUT_DIR}\\{Manager_ID}_unanswered.json
         {OUTPUT_DIR}\\{Manager_ID}_unanswered.docx
    - Always prompts for next question until user types 'exit' or 'quit'.

Config:
- INPUT_DIR / OUTPUT_DIR are set as constants below (edit in-code later).
- Requires OPENAI_API_KEY env var.

Install:
    pip install openai python-docx scikit-learn

Run:
    python aggregate_qa_llm_v1.py --manager_id 10067 --user_id 12345
"""

from __future__ import annotations

import argparse
import json
import os
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

# ---------- EDIT THESE DEFAULTS IN CODE ----------
INPUT_DIR = r"C:\Users\prash\Downloads"
OUTPUT_DIR = r"C:\Users\prash\Downloads"

DEFAULT_MODEL = "gpt-4o-2024-08-06"
TOP_K = 40  # bullets passed to LLM
# -----------------------------------------------

SOURCE_PRIORITY = {
    "Townsend": 3,
    "3rd party": 2,
    "Third party": 2,
    "Third-party": 2,
    "Manager": 1,
    "Unknown": 0,
}

def normalize_text(s: str) -> str:
    s = (s or "").lower()
    s = re.sub(r"\s+", " ", s).strip()
    return s

def source_priority(src: str) -> int:
    if not src:
        return 0
    return SOURCE_PRIORITY.get(src, SOURCE_PRIORITY.get(src.title(), 0))

def parse_date(s: str) -> Optional[datetime]:
    if not s:
        return None
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S%z"):
        try:
            return datetime.strptime(s, fmt)
        except Exception:
            pass
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00"))
    except Exception:
        return None

def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00","Z")

def load_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def save_json(path: str, obj: Any) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2, ensure_ascii=False)

# ---------- Retrieval ----------

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    _SKLEARN_OK = True
except Exception:
    _SKLEARN_OK = False

def flatten_aggregate(agg: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    sections = agg.get("sections", {}) or {}
    for section_id, sdata in sections.items():
        bullets = (sdata or {}).get("bullets", []) or []
        for b in bullets:
            prov = b.get("provenance", {}) or {}
            rows.append({
                "bullet_uid": b.get("bullet_uid"),
                "section_id": section_id,
                "text": b.get("text", "") or "",
                "text_norm": b.get("text_norm") or normalize_text(b.get("text", "")),
                "score": float(b.get("score", 0.0) or 0.0),
                "document_id": prov.get("document_id"),
                "document_source": prov.get("document_source", "Unknown"),
                "document_date": prov.get("document_date"),
                "provenance": prov,
            })
    return rows

def build_tfidf_index(rows: List[Dict[str, Any]]):
    corpus = [r["text_norm"] for r in rows]
    vec = TfidfVectorizer(min_df=1, max_df=0.95, ngram_range=(1, 2))
    X = vec.fit_transform(corpus)
    return vec, X

def rerank(rows: List[Dict[str, Any]], base_scores: List[float]) -> List[Tuple[float, Dict[str, Any]]]:
    scored: List[Tuple[float, Dict[str, Any]]] = []
    for r, s in zip(rows, base_scores):
        dd = parse_date(r.get("document_date") or "")
        recency = 0.0
        if dd:
            # Coerce naive datetimes to UTC to avoid offset-naive/aware subtraction
            if dd.tzinfo is None:
                dd = dd.replace(tzinfo=timezone.utc)
            days = (datetime.now(timezone.utc) - dd).days
            recency = max(0.0, 1.0 - min(days, 365*5) / (365*5))
        sp = source_priority(r.get("document_source"))
        final = (float(s) * 1.0) + (recency * 0.15) + (sp * 0.05) + (float(r.get("score", 0.0)) * 0.02)
        scored.append((final, r))
    scored.sort(key=lambda x: x[0], reverse=True)
    return scored

def retrieve(rows: List[Dict[str, Any]], query: str, top_k: int = TOP_K) -> List[Dict[str, Any]]:
    query_norm = normalize_text(query)
    if not rows:
        return []
    if _SKLEARN_OK:
        vec, X = build_tfidf_index(rows)
        qv = vec.transform([query_norm])
        sims = (X @ qv.T).toarray().flatten()
        ranked = rerank(rows, sims.tolist())
        return [r for _, r in ranked[:top_k]]
    # Fallback: token overlap
    qset = set(query_norm.split())
    scored = []
    for r in rows:
        tset = set((r["text_norm"] or "").split())
        inter = len(qset & tset)
        union = len(qset | tset) or 1
        scored.append(inter / union)
    ranked = rerank(rows, scored)
    return [r for _, r in ranked[:top_k]]

# ---------- LLM call (Structured Outputs via json_schema) ----------

def call_llm_grounded_answer(
    model: str,
    user_id: str,
    manager_id: str,
    question: str,
    evidence: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    Returns dict with keys:
      - sufficient (bool)
      - answer (str)
      - confidence (HIGH|MEDIUM|LOW)
      - citations (list[str])
      - insufficiency_reason (str|null)
      - follow_up_questions (list[str])
    """
    from openai import OpenAI
    client = OpenAI()

    ev_payload = []
    for r in evidence:
        prov = r.get("provenance") or {}
        ev_payload.append({
            "bullet_uid": r.get("bullet_uid"),
            "section_id": r.get("section_id"),
            "text": r.get("text"),
            "document_source": prov.get("document_source", r.get("document_source")),
            "document_date": prov.get("document_date", r.get("document_date")),
            "document_id": prov.get("document_id", r.get("document_id")),
        })

    schema = {
        "name": "grounded_manager_answer",
        "strict": True,
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "sufficient": {"type": "boolean"},
                "confidence": {"type": "string", "enum": ["HIGH", "MEDIUM", "LOW"]},
                "answer": {"type": "string"},
                "citations": {"type": "array", "items": {"type": "string"}},
                "insufficiency_reason": {"type": ["string", "null"]},
                "follow_up_questions": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["sufficient", "confidence", "answer", "citations", "insufficiency_reason", "follow_up_questions"]
        }
    }

    system = (
        "You are a diligence analyst answering questions about an investment manager.\n"
        "You MUST use ONLY the provided evidence bullets.\n"
        "If the evidence is insufficient to answer the question, set sufficient=false, "
        "explain why in insufficiency_reason, and provide 2-5 follow_up_questions.\n"
        "If sufficient=true, write a professional, organized answer (headings/bullets encouraged), "
        "and include citations as a list of bullet_uid values you relied on.\n"
        "Never guess. Never introduce facts not present in the evidence.\n"
        "If evidence is conflicting, explicitly note the conflict and prefer newer document_date and "
        "Townsend source where applicable, while acknowledging alternatives."
    )

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps({
                "manager_id": manager_id,
                "question": question,
                "evidence_bullets": ev_payload
            })}
        ],
        response_format={"type": "json_schema", "json_schema": schema},
        safety_identifier=user_id,
        temperature=0.2,
    )

    content = resp.choices[0].message.content
    if not content:
        return {
            "sufficient": False,
            "confidence": "LOW",
            "answer": "",
            "citations": [],
            "insufficiency_reason": "Model returned empty response.",
            "follow_up_questions": ["Try rephrasing the question."],
        }

    try:
        return json.loads(content)
    except Exception:
        return {
            "sufficient": False,
            "confidence": "LOW",
            "answer": "",
            "citations": [],
            "insufficiency_reason": "Failed to parse model output as JSON.",
            "follow_up_questions": ["Try again or reduce evidence size."],
        }

# ---------- Unanswered logging ----------

def unanswered_paths(manager_id: str) -> Tuple[str, str]:
    return (
        os.path.join(OUTPUT_DIR, f"{manager_id}_unanswered.json"),
        os.path.join(OUTPUT_DIR, f"{manager_id}_unanswered.docx"),
    )

def append_unanswered_json(manager_id: str, record: Dict[str, Any]) -> None:
    json_path, _ = unanswered_paths(manager_id)
    if os.path.exists(json_path):
        try:
            arr = load_json(json_path)
            if not isinstance(arr, list):
                arr = []
        except Exception:
            arr = []
    else:
        arr = []
    arr.append(record)
    save_json(json_path, arr)

def append_unanswered_docx(manager_id: str, record: Dict[str, Any]) -> None:
    _, docx_path = unanswered_paths(manager_id)
    from docx import Document

    if os.path.exists(docx_path):
        doc = Document(docx_path)
    else:
        doc = Document()
        doc.add_heading(f"Unanswered / Unsatisfactory Questions — Manager {manager_id}", level=1)

    doc.add_paragraph(f"Asked at: {record.get('asked_at')}")
    doc.add_paragraph(f"User_ID: {record.get('user_id')}")
    doc.add_paragraph(f"Status: {record.get('status')}")
    doc.add_paragraph(f"Question: {record.get('question')}")

    reason = record.get("insufficiency_reason") or record.get("notes")
    if reason:
        p = doc.add_paragraph("Reason / Notes:")
        p.runs[0].bold = True
        doc.add_paragraph(str(reason))

    fu = record.get("follow_up_questions") or []
    if fu:
        p = doc.add_paragraph("Follow-up questions:")
        p.runs[0].bold = True
        for q in fu[:10]:
            doc.add_paragraph(f"- {q}")

    cited = record.get("citations") or []
    if cited:
        p = doc.add_paragraph("Citations (bullet_uid):")
        p.runs[0].bold = True
        doc.add_paragraph(", ".join(cited[:50]))

    cand = (record.get("retrieval") or {}).get("candidate_bullet_uids") or []
    if cand:
        p = doc.add_paragraph("Retrieved candidates (bullet_uid, first 25):")
        p.runs[0].bold = True
        doc.add_paragraph(", ".join(cand[:25]))

    doc.add_page_break()
    doc.save(docx_path)

# ---------- Main loop ----------

def main():
    ap = argparse.ArgumentParser(description="Grounded Q&A over {manager_id}_aggregate.json (LLM-assisted).")
    ap.add_argument("--manager_id", required=True, help="Manager_ID (e.g., 10067)")
    ap.add_argument("--user_id", required=True, help="User_ID in xxxxx format (recorded in logs)")
    ap.add_argument("--model", default=DEFAULT_MODEL, help=f"OpenAI model (default: {DEFAULT_MODEL})")
    args = ap.parse_args()

    manager_id = args.manager_id.strip()
    user_id = args.user_id.strip()
    model = args.model.strip()

    agg_path = os.path.join(INPUT_DIR, f"{manager_id}_aggregate.json")
    if not os.path.exists(agg_path):
        raise FileNotFoundError(f"Aggregate not found: {agg_path}")

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    agg = load_json(agg_path)
    rows = flatten_aggregate(agg)
    manager_name = agg.get("manager_name") or manager_id

    print(f"\nLoaded manager {manager_name} ({manager_id})")
    print(f"Aggregate: {agg_path}")
    print(f"Output dir: {OUTPUT_DIR}")
    print("Type 'exit' or 'quit' to stop.\n")

    while True:
        q = input("Question> ").strip()
        if not q:
            continue
        if q.lower() in ("exit", "quit"):
            break

        candidates = retrieve(rows, q, top_k=TOP_K)
        candidate_uids = [c.get("bullet_uid") for c in candidates if c.get("bullet_uid")]

        result = call_llm_grounded_answer(
            model=model,
            user_id=user_id,
            manager_id=manager_id,
            question=q,
            evidence=candidates
        )

        sufficient = bool(result.get("sufficient"))
        answer = (result.get("answer") or "").strip()
        confidence = result.get("confidence") or "LOW"
        citations = result.get("citations") or []
        insufficiency_reason = result.get("insufficiency_reason")
        follow_up = result.get("follow_up_questions") or []

        print("\n--- Answer ---")
        if sufficient and answer:
            print(answer)
        else:
            print("INSUFFICIENT INFORMATION to answer from the aggregate evidence.")
            if insufficiency_reason:
                print(f"\nReason: {insufficiency_reason}")
            if follow_up:
                print("\nSuggested follow-ups:")
                for fq in follow_up[:5]:
                    print(f"- {fq}")
        print("--------------\n")

        status = None
        notes = None
        thumbs_down = False

        if sufficient and answer:
            fb = input("Satisfied? (y/n) > ").strip().lower()
            if fb in ("n", "no"):
                thumbs_down = True
                status = "UNSATISFACTORY_ANSWER"
                notes = "User indicated thumbs-down."
            else:
                status = "ANSWERED_SATISFACTORY"
        else:
            status = "INSUFFICIENT_INFORMATION"

        if (not sufficient) or thumbs_down:
            record = {
                "asked_at": utc_now_iso(),
                "manager_id": manager_id,
                "manager_name": manager_name,
                "user_id": user_id,
                "question": q,
                "status": status,
                "confidence": confidence,
                "citations": citations,
                "insufficiency_reason": insufficiency_reason,
                "follow_up_questions": follow_up,
                "notes": notes,
                "retrieval": {
                    "top_k": TOP_K,
                    "candidate_bullet_uids": candidate_uids,
                },
            }
            append_unanswered_json(manager_id, record)
            append_unanswered_docx(manager_id, record)
            print(f"(Logged to unanswered files for manager {manager_id}.)\n")

    print("Done.")

if __name__ == "__main__":
    main()
