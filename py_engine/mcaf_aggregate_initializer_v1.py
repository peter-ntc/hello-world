#!/usr/bin/env python3
"""
MCAF Aggregate Initializer v1

Creates empty per-manager aggregate artifacts:
  - <manager_id>_aggregate.json
  - <manager_id>_aggregate.docx (optional; default ON)

Design goals:
  - Deterministic, stable empty structure with S01–S15 sections
  - Default output directory: C:/Users/prash/Downloads/ (if exists), else current working directory

Note: We use forward slashes in this docstring to avoid Python's backslash escape sequences.
"""

from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document


# ---------- MCAF SECTION DEFINITIONS (top-level only) ----------

MCAF_SECTIONS: List[Tuple[str, str]] = [
    ("S01", "Team Experience & Structure"),
    ("S02", "Organizational Depth & Infrastructure"),
    ("S03", "Sector-Specific Capabilities"),
    ("S04", "Investment Process & Governance"),
    ("S05", "Alignment of Interest"),
    ("S06", "Fiduciary Standards & Reputation"),
    ("S07", "Track Record Quality"),
    ("S08", "Platform Evolution & Stability"),
    ("S09", "Vertical Integration & Operating Edge"),
    ("S10", "Conflicts & Transparency"),
    ("S11", "Institutional Readiness"),
    ("S12", "Fee & Incentive Assessment"),
    ("S13", "Team Culture & Retention"),
    ("S14", "Overall Manager Rating"),
    ("S15", "Other Comprehensive Manager Information"),
]

DEFAULT_OUTPUT_DIR = r"C:\Users\prash\Downloads"


def parse_yes_no(val: str) -> bool:
    v = (val or "").strip().lower()
    if v in ("y", "yes", "true", "1", "on"):
        return True
    if v in ("n", "no", "false", "0", "off"):
        return False
    raise argparse.ArgumentTypeError("DOCX_WRITE must be Yes/No (or True/False).")


def resolve_output_dir(arg_output_dir: str | None) -> Path:
    if arg_output_dir:
        out = Path(arg_output_dir).expanduser().resolve()
        out.mkdir(parents=True, exist_ok=True)
        return out

    default = Path(DEFAULT_OUTPUT_DIR)
    if default.exists() and default.is_dir():
        return default.resolve()

    # fallback: current working directory
    return Path.cwd().resolve()


def build_empty_aggregate_json(manager_id: int) -> Dict:
    run_dt = datetime.now(timezone.utc).replace(microsecond=0).isoformat()

    sections = {
        sid: {"section_id": sid, "section_title": title, "bullets": []}
        for sid, title in MCAF_SECTIONS
    }

    return {
        "schema_version": "mcaf_aggregate_v1",
        "manager_id": manager_id,
        "manager_name": "",  # can be populated later when first MCAF doc is merged
        "created_datetime": run_dt,
        "documents": [],
        "sections": sections,
        "coverage": {
            "document_count": 0,
            "bullet_occurrence_count": 0,
            "unique_bullet_uid_count": 0,
        },
        "notes": {
            "docx_write_default": True,
            "purpose": "Append-only aggregate store for all per-document MCAF filings for this manager.",
        },
    }


def write_empty_docx(manager_id: int, output_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"MCAF Aggregate (Empty): Manager {manager_id}", level=0)
    doc.add_paragraph(
        "This document is an aggregate shell. It will be populated by the aggregator as MCAF filings are merged."
    )
    doc.add_paragraph(f"Manager ID: {manager_id}")

    for sid, title in MCAF_SECTIONS:
        doc.add_heading(f"{sid}: {title}", level=1)
        doc.add_paragraph("(empty)")

    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Initialize empty per-manager MCAF aggregate JSON/DOCX.")
    parser.add_argument("--manager_id", required=True, type=int, help="Manager ID (e.g., 10067)")
    parser.add_argument(
        "--output_dir",
        default=None,
        help=r"Output directory (default: C:\Users\prash\Downloads if it exists, else current directory)",
    )
    parser.add_argument(
        "--DOCX_WRITE",
        type=parse_yes_no,
        default=True,
        help="Whether to write the DOCX shell (Yes/No). Default: Yes",
    )
    args = parser.parse_args()

    out_dir = resolve_output_dir(args.output_dir)
    manager_id = int(args.manager_id)

    base = f"{manager_id}_aggregate"
    out_json = out_dir / f"{base}.json"
    out_docx = out_dir / f"{base}.docx"

    agg = build_empty_aggregate_json(manager_id)

    # Write JSON (always)
    with out_json.open("w", encoding="utf-8") as f:
        json.dump(agg, f, indent=2, ensure_ascii=False)

    # Write DOCX (optional)
    if args.DOCX_WRITE:
        write_empty_docx(manager_id, out_docx)

    print("Wrote:")
    print(f"  {out_json}")
    if args.DOCX_WRITE:
        print(f"  {out_docx}")
    else:
        print("  (DOCX_WRITE=No; DOCX not written)")


if __name__ == "__main__":
    main()
