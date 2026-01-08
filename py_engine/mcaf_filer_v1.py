#!/usr/bin/env python3
"""
MCAF Filer v1

Input: RAW JSON produced by manager_views_extractor v7 (single document).
Output:
  - <manager_id>_mcaf_<document_id>.json
  - <manager_id>_mcaf_<document_id>.docx
  - <manager_id>_mcaf_<document_id>_audit.json

Key rules enforced:
  - File into 15 top-level sections (S01–S15) only (no subsections)
  - Each bullet assigned to 1–3 sections; if >3 suggested, keep top 3 by score
  - Duplication across sections allowed; within-section duplicates disallowed (document-aware)
  - Provenance preserved per bullet (doc/date/source/extraction_datetime)
  - DOCX ordering within each section: newest date first; within date Townsend > Manager > 3rd party
"""

from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional

from docx import Document

# ---------- MCAF SECTION DEFINITIONS (top-level only) ----------

MCAF_SECTIONS: List[Tuple[str, str]] = [
    ("S01", "Team Experience & Structure"),
    ("S02", "Organizational Depth & Infrastructure"),
    ("S03", "Sector-Specific Capabilities"),
    ("S04", "Investment Process & Governance"),
    ("S05", "Alignment of Interest"),
    ("S06", "Fiduciary Standards & Reputation"),
    ("S07", "Track Record Quality"),
    ("S08", "Platform Evolution & Stability"),
    ("S09", "Vertical Integration & Operating Edge"),
    ("S10", "Conflicts & Transparency"),
    ("S11", "Institutional Readiness"),
    ("S12", "Fee & Incentive Assessment"),
    ("S13", "Team Culture & Retention"),
    ("S14", "Overall Manager Rating"),
    ("S15", "Other Comprehensive Manager Information"),
]

SOURCE_PRIORITY = {
    "Townsend": 0,
    "Manager": 1,
    "3rd party": 2,
}

AUDIT_HEADER_PREFIXES = (
    "Document Date:",
    "Extraction Date:",
    "Document Source:",
)

SECTION_HEADER_PATTERNS = [
    re.compile(r"^\s*##\s*"),              # markdown-ish headings like "## 1) ..."
    re.compile(r"^\s*\d+\)\s+"),           # "1) Identity ..."
    re.compile(r"^\s*#+\s*"),              # any markdown heading
]

BULLET_RE = re.compile(r"^(\s*)-\s+(.*)$")  # captures indent + text


# ---------- DATA MODELS ----------

@dataclass(frozen=True)
class Provenance:
    manager_id: int
    manager_name: str
    document_id: int
    document_date: str  # YYYY-MM-DD
    document_source: str  # Townsend / Manager / 3rd party
    extraction_datetime: str  # ISO

@dataclass
class AtomicBullet:
    raw_id: str
    text: str
    text_norm: str
    order: int
    provenance: Provenance


# ---------- UTILITIES ----------

def normalize_text(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = s.rstrip(" .;:,")
    return s

def is_section_header(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    # Avoid treating bullets as headers
    if stripped.startswith("-"):
        return False
    return any(p.search(stripped) for p in SECTION_HEADER_PATTERNS)

def parse_raw_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)

def safe_parse_date(d: str) -> datetime:
    # expects YYYY-MM-DD
    return datetime.strptime(d, "%Y-%m-%d")

def ensure_source_canonical(src: str) -> str:
    # extractor v7 uses canonical strings; keep defensive normalization
    s = (src or "").strip().lower()
    if s in ("t", "townsend"):
        return "Townsend"
    if s in ("m", "manager"):
        return "Manager"
    if s in ("3", "3rd", "3rd party", "third party", "third-party"):
        return "3rd party"
    # fallback: title-case but keep 3rd party exactly if present
    if "3" in s and "party" in s:
        return "3rd party"
    return src.strip() or "3rd party"


# ---------- PARSING: profile_text -> atomic bullets (deterministic, no LLM) ----------

def _strip_audit_headers(lines: List[str]) -> List[str]:
    out = []
    for ln in lines:
        if any(ln.strip().startswith(p) for p in AUDIT_HEADER_PREFIXES):
            continue
        out.append(ln)
    return out

def extract_atomic_bullets(profile_text: str, prov: Provenance) -> List[AtomicBullet]:
    """
    Deterministic extraction of atomic bullets from profile_text, preserving indentation.
    Rules:
      - Remove audit header lines.
      - Ignore section headers (e.g., '## 1) ...').
      - Parse lines starting with '-' as bullets.
      - Collapse "container bullets" that end with ':' by appending subsequent bullet lines
        until next container or header.
      - Attach indented sub-bullets to their parent bullet (flattened into one atomic bullet).
    """
    raw_lines = [ln.rstrip("\n") for ln in profile_text.splitlines()]
    raw_lines = [ln for ln in raw_lines if ln.strip()]
    raw_lines = _strip_audit_headers(raw_lines)

    bullets: List[AtomicBullet] = []
    i = 0
    order = 0

    while i < len(raw_lines):
        ln = raw_lines[i]
        if is_section_header(ln):
            i += 1
            continue

        m = BULLET_RE.match(ln)
        if not m:
            i += 1
            continue

        indent = len(m.group(1).replace("\t", "    "))
        text = m.group(2).strip()

        # Gather potential sub-bullets (indented more than current indent)
        def peek_bullet(idx: int):
            if idx >= len(raw_lines):
                return None
            mm = BULLET_RE.match(raw_lines[idx])
            if not mm:
                return None
            ind = len(mm.group(1).replace("\t", "    "))
            return ind, mm.group(2).strip()

        # Container bullet logic (endswith ':' and has content)
        if text.endswith(":") and len(text) > 1:
            container = text[:-1].strip()
            subs: List[str] = []
            j = i + 1
            while j < len(raw_lines):
                nxt = raw_lines[j]
                if is_section_header(nxt):
                    break
                pk = peek_bullet(j)
                if pk is None:
                    j += 1
                    continue
                n_indent, n_text = pk
                # Stop if next is another container at same indent (or less)
                if n_text.endswith(":") and n_indent <= indent:
                    break
                # Accept bullets that are deeper or same-level; we flatten all as claims under container
                subs.append(n_text)
                j += 1
            combined = container + ": " + "; ".join(subs) if subs else container + ":"
            order += 1
            raw_id = f"R{order:03d}"
            bullets.append(
                AtomicBullet(
                    raw_id=raw_id,
                    text=combined,
                    text_norm=normalize_text(combined),
                    order=order,
                    provenance=prov,
                )
            )
            i = j
            continue

        # Non-container bullet: attach indented sub-bullets directly beneath it
        subs: List[str] = []
        j = i + 1
        while j < len(raw_lines):
            nxt = raw_lines[j]
            if is_section_header(nxt):
                break
            pk = peek_bullet(j)
            if pk is None:
                j += 1
                continue
            n_indent, n_text = pk
            # Stop at same-or-less indent (new sibling/parent)
            if n_indent <= indent:
                break
            # If an indented container appears, treat it as just another sub-item text (keep ':')
            subs.append(n_text)
            j += 1

        combined = text
        if subs:
            combined = combined + " — " + "; ".join(subs)

        order += 1
        raw_id = f"R{order:03d}"
        bullets.append(
            AtomicBullet(
                raw_id=raw_id,
                text=combined,
                text_norm=normalize_text(combined),
                order=order,
                provenance=prov,
            )
        )
        i = j

    return bullets


# ---------- CLASSIFICATION (LLM): bullet -> up to 3 MCAF sections ----------

def classify_bullets_openai(
    bullets: List[AtomicBullet],
    model: str = "gpt-5.2",
    temperature: float = 0.1,
    batch_size: int = 25,
) -> Dict[str, List[Dict[str, float]]]:
    """
    Returns mapping: raw_id -> list of assignments [{section_id, score}, ...] sorted desc.
    Requires OPENAI_API_KEY in environment.
    """
    try:
        from openai import OpenAI
    except Exception as e:
        raise RuntimeError(
            "openai package not available. Install with: pip install openai"
        ) from e

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set in environment.")

    client = OpenAI(api_key=api_key)

    sections_payload = [{"section_id": sid, "title": title} for sid, title in MCAF_SECTIONS]

    system = (
        "You are a classification engine. You must only classify each bullet into the most relevant MCAF sections.\n"
        "Do not rewrite or summarize bullets. Do not invent facts. Output MUST be valid JSON only."
    )

    results: Dict[str, List[Dict[str, float]]] = {}

    for start in range(0, len(bullets), batch_size):
        batch = bullets[start : start + batch_size]
        bullets_payload = [{"raw_id": b.raw_id, "text": b.text} for b in batch]

        user = {
            "task": "Classify each bullet into 1 to 3 MCAF sections.",
            "rules": [
                "Return between 1 and 3 section assignments per bullet.",
                "If more than 3 could apply, choose ONLY the top 3 most relevant.",
                "Scores must be floats between 0 and 1.",
                "Assignments must be sorted by descending score.",
                "Never output invalid section IDs.",
                "Output JSON only. No prose.",
            ],
            "mcaf_sections": sections_payload,
            "bullets": bullets_payload,
            "output_schema": {
                "results": [
                    {
                        "raw_id": "R001",
                        "assignments": [
                            {"section_id": "S07", "score": 0.92},
                            {"section_id": "S04", "score": 0.71},
                        ],
                    }
                ]
            },
        }

        resp = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": json.dumps(user)},
            ],
            response_format={"type": "json_object"},
        )

        content = resp.choices[0].message.content
        parsed = json.loads(content)

        for item in parsed.get("results", []):
            rid = item.get("raw_id")
            assigns = item.get("assignments") or []
            cleaned = []
            seen = set()
            for a in assigns:
                sid = a.get("section_id")
                score = float(a.get("score", 0.0))
                if sid not in dict(MCAF_SECTIONS):
                    continue
                if sid in seen:
                    continue
                seen.add(sid)
                cleaned.append({"section_id": sid, "score": score})
            cleaned.sort(key=lambda x: x["score"], reverse=True)
            results[rid] = cleaned

    return results


# ---------- ENFORCEMENT + OUTPUT JSON BUILD ----------

def build_mcaf_json(
    bullets: List[AtomicBullet],
    assignments: Dict[str, List[Dict[str, float]]],
    run_datetime: str,
    model: str,
    prompt_version: str = "mcaf_filer_v1",
) -> Tuple[dict, dict]:
    """
    Returns (mcaf_json, audit_json)
    """
    section_titles = dict(MCAF_SECTIONS)

    # Init section containers
    sections: Dict[str, dict] = {
        sid: {
            "section_id": sid,
            "section_title": section_titles[sid],
            "bullets": [],
        }
        for sid, _ in MCAF_SECTIONS
    }

    forced_fallback_ids: List[str] = []
    assigned_over_3: List[str] = []
    missing_assignments: List[str] = []
    duplicates_within_section: Dict[str, List[dict]] = {}

    # Within-section dedupe sets: document-aware
    # For same document, dedupe by normalized text; across documents, allow.
    dedupe_seen: Dict[str, set] = {sid: set() for sid, _ in MCAF_SECTIONS}

    multi_assigned_counts = {1: 0, 2: 0, 3: 0}

    for b in bullets:
        cand = assignments.get(b.raw_id, [])
        if not cand:
            # ensure coverage; force S15
            cand = [{"section_id": "S15", "score": 0.01}]
            forced_fallback_ids.append(b.raw_id)

        # enforce max 3
        cand = sorted(cand, key=lambda x: x["score"], reverse=True)
        if len(cand) > 3:
            assigned_over_3.append(b.raw_id)
            cand = cand[:3]

        # ensure at least 1
        if len(cand) == 0:
            missing_assignments.append(b.raw_id)
            cand = [{"section_id": "S15", "score": 0.01}]
            forced_fallback_ids.append(b.raw_id)

        multi_assigned_counts[len(cand)] = multi_assigned_counts.get(len(cand), 0) + 1

        for rank_idx, a in enumerate(cand, start=1):
            sid = a["section_id"]
            score = float(a.get("score", 0.0))

            # document-aware within-section dedupe:
            # duplicates are only duplicates if same document_id AND same normalized text
            key = (b.provenance.document_id, b.text_norm)
            if key in dedupe_seen[sid]:
                duplicates_within_section.setdefault(sid, []).append(
                    {"raw_id": b.raw_id, "document_id": b.provenance.document_id, "text": b.text}
                )
                continue
            dedupe_seen[sid].add(key)

            sections[sid]["bullets"].append(
                {
                    "raw_id": b.raw_id,
                    "text": b.text,
                    "provenance": {
                        "document_id": b.provenance.document_id,
                        "document_date": b.provenance.document_date,
                        "document_source": b.provenance.document_source,
                        "extraction_datetime": b.provenance.extraction_datetime,
                    },
                    "rank": rank_idx,
                    "score": score,
                    "assigned_by": {
                        "model": model,
                        "prompt_version": prompt_version,
                        "run_datetime": run_datetime,
                    },
                }
            )

    # Coverage check: every raw_id must appear in >=1 section (after dedupe, still should)
    all_raw_ids = {b.raw_id for b in bullets}
    filed_raw_ids = set()
    for sid in sections:
        for item in sections[sid]["bullets"]:
            filed_raw_ids.add(item["raw_id"])

    missing_raw_ids = sorted(list(all_raw_ids - filed_raw_ids))

    mcaf_json = {
        "schema_version": "mcaf_filer_v1",
        "manager_id": bullets[0].provenance.manager_id if bullets else None,
        "manager_name": bullets[0].provenance.manager_name if bullets else None,
        "documents": [
            {
                "document_id": bullets[0].provenance.document_id if bullets else None,
                "document_date": bullets[0].provenance.document_date if bullets else None,
                "document_source": bullets[0].provenance.document_source if bullets else None,
                "extraction_datetime": bullets[0].provenance.extraction_datetime if bullets else None,
            }
        ],
        "sections": sections,
        "coverage": {
            "raw_bullet_count": len(bullets),
            "filed_unique_raw_ids": len(filed_raw_ids),
            "missing_raw_ids": missing_raw_ids,
            "multi_assigned_counts": multi_assigned_counts,
        },
    }

    audit_json = {
        "run_datetime": run_datetime,
        "forced_fallback_ids": forced_fallback_ids,
        "assigned_over_3_ids": assigned_over_3,  # should be empty once enforced (kept for visibility)
        "duplicates_within_section": duplicates_within_section,  # should be empty
        "missing_raw_ids": missing_raw_ids,  # should be empty
        "missing_assignments_pre_fallback": missing_assignments,
    }

    return mcaf_json, audit_json


# ---------- DOCX RENDERING (human consumption; provenance-driven ordering) ----------

def render_mcaf_docx(mcaf_json: dict, output_path: Path) -> None:
    doc = Document()

    manager_name = mcaf_json.get("manager_name", "")
    manager_id = mcaf_json.get("manager_id", "")
    doc_meta = (mcaf_json.get("documents") or [{}])[0]
    document_id = doc_meta.get("document_id", "")
    document_date = doc_meta.get("document_date", "")
    document_source = doc_meta.get("document_source", "")

    doc.add_heading(f"MCAF Profile: {manager_name} (ID {manager_id})", level=0)
    doc.add_paragraph(f"Input Document ID: {document_id}")
    doc.add_paragraph(f"Document Date: {document_date}")
    doc.add_paragraph(f"Document Source: {document_source}")
    doc.add_paragraph(f"Filer Run Datetime: {mcaf_json.get('coverage', {}).get('run_datetime', '')}".strip())

    sections = mcaf_json["sections"]

    # Render in fixed S01..S15 order
    for sid, title in MCAF_SECTIONS:
        doc.add_heading(f"{sid}: {title}", level=1)

        items = sections[sid]["bullets"]

        # Sort by date desc, then source priority, then raw_id
        def sort_key(item):
            prov = item["provenance"]
            d = prov.get("document_date", "1900-01-01")
            try:
                dt = safe_parse_date(d)
            except Exception:
                dt = datetime(1900, 1, 1)
            src = prov.get("document_source", "3rd party")
            src = ensure_source_canonical(src)
            return (-int(dt.timestamp()), SOURCE_PRIORITY.get(src, 99), item.get("raw_id", ""))

        items_sorted = sorted(items, key=sort_key)

        for it in items_sorted:
            prov = it["provenance"]
            d = prov.get("document_date", "")
            src = ensure_source_canonical(prov.get("document_source", ""))
            docid = prov.get("document_id", "")
            rid = it.get("raw_id", "")
            text = it.get("text", "")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
            p.add_run(f"  [{d} | {src} | Doc {docid} | {rid}]").italic = True

    doc.save(str(output_path))


# ---------- MAIN ----------

def main():
    parser = argparse.ArgumentParser(description="MCAF Filer v1 (RAW JSON -> MCAF JSON + DOCX + audit)")
    parser.add_argument("--input", required=True, help="Path to RAW JSON from extractor v7")
    parser.add_argument("--output_dir", default=None, help="Output directory (default: same directory as input)")
    parser.add_argument("--model", default="gpt-5.2", help="OpenAI model for classification (default: gpt-5.2)")
    parser.add_argument("--temperature", type=float, default=0.1, help="Classification temperature (default: 0.1)")
    parser.add_argument("--batch_size", type=int, default=25, help="Bullets per classification call (default: 25)")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    raw = parse_raw_json(input_path)

    manager_id = int(raw.get("manager_id"))
    manager_name = raw.get("manager_name", "")
    doc_meta = raw.get("document_metadata") or {}
    document_id = int(doc_meta.get("document_id"))
    document_date = doc_meta.get("document_date")
    document_source = ensure_source_canonical(doc_meta.get("document_source", ""))
    extraction_datetime = doc_meta.get("extraction_datetime") or (raw.get("document_metadata", {}) or {}).get("extraction_datetime")

    profile_text = raw.get("profile_text")
    if not profile_text:
        raise ValueError("RAW JSON missing profile_text")

    prov = Provenance(
        manager_id=manager_id,
        manager_name=manager_name,
        document_id=document_id,
        document_date=document_date,
        document_source=document_source,
        extraction_datetime=extraction_datetime,
    )

    bullets = extract_atomic_bullets(profile_text=profile_text, prov=prov)
    if not bullets:
        raise ValueError("No bullets parsed from profile_text. Check RAW format or parser rules.")

    run_dt = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

    assignments = classify_bullets_openai(
        bullets=bullets,
        model=args.model,
        temperature=args.temperature,
        batch_size=args.batch_size,
    )

    mcaf_json, audit_json = build_mcaf_json(
        bullets=bullets,
        assignments=assignments,
        run_datetime=run_dt,
        model=args.model,
    )
    # Store run datetime in coverage for DOCX header convenience
    mcaf_json.setdefault("coverage", {})["run_datetime"] = run_dt

    out_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else input_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    base = f"{manager_id}_mcaf_{document_id}"
    out_json = out_dir / f"{base}.json"
    out_docx = out_dir / f"{base}.docx"
    out_audit = out_dir / f"{base}_audit.json"

    with out_json.open("w", encoding="utf-8") as f:
        json.dump(mcaf_json, f, indent=2, ensure_ascii=False)

    with out_audit.open("w", encoding="utf-8") as f:
        json.dump(audit_json, f, indent=2, ensure_ascii=False)

    render_mcaf_docx(mcaf_json, out_docx)

    print("Wrote:")
    print(f"  {out_json}")
    print(f"  {out_docx}")
    print(f"  {out_audit}")


if __name__ == "__main__":
    main()
