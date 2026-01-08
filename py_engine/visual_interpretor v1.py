#!/usr/bin/env python3
r"""
visual_interpretor.py

Batch-interpret extracted table/graph images using GPT (vision) and append results to:
  - <MANAGER_ID>_rawvisual_<DOCUMENT_ID>.json
  - <MANAGER_ID>_rawvisual_<DOCUMENT_ID>.docx

Design goals (per your spec):
- Inventory directory holds ALL visuals: <output_root>/<MANAGER_ID>_visuals/
- Interpreter creates/uses a holding subdir: <visuals_dir>/batch/
- Each run selects the next N unprocessed visuals (ledger = existing rawvisual JSON),
  copies them into batch/, sends them to GPT in ONE request, then appends outputs.
- N is configurable via --batch_size with a hard cap of 15:
    * default = 15
    * if specified > 15 -> treated as 15
- Visuals are referenced in JSON/DOCX so they can be retrieved alongside insights.
- PDF-only for now, but CLI mirrors your pipeline metadata.

Dependencies:
  pip install openai python-docx

Environment:
  Set OPENAI_API_KEY in your shell before running:
    Windows CMD: set OPENAI_API_KEY=YOUR_KEY

Example:
  python visual_interpretor.py ^
    --input "C:/path/to/deck.pdf" ^
    --manager "Asana" ^
    --manager_id 10067 ^
    --document_id 55666 ^
    --document_date 2024-11-15 ^
    --document_source "Manager" ^
    --visuals_dir "C:/Users/prash/Downloads/10067_visuals" ^
    --batch_size 15
"""

from __future__ import annotations

import argparse
import base64
import json
import os
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document
from openai import OpenAI


MAX_BATCH = 15
DOCID_FILE_RE = re.compile(r"^(?P<docid>\d+)_(?P<num>\d{4})\.png$", re.IGNORECASE)


# -----------------------------
# Safety validations (manager_id / document_id consistency)
# -----------------------------
def validate_visuals_dir_manager_id(visuals_dir: Path, manager_id: int) -> None:
    """
    If visuals_dir basename matches '<digits>_visuals', enforce that digits == manager_id.
    Prevents accidentally pointing at another manager's inventory directory.
    """
    name = visuals_dir.name
    m = re.match(r"^(?P<mid>\d+)_visuals$", name, flags=re.IGNORECASE)
    if m:
        dir_mid = int(m.group("mid"))
        if dir_mid != manager_id:
            raise ValueError(
                f"--visuals_dir looks like '{name}' (manager_id={dir_mid}) but CLI --manager_id={manager_id}. "
                f"Point --visuals_dir to '{manager_id}_visuals' or fix --manager_id."
            )


def present_document_ids(visuals_dir: Path) -> list[int]:
    """
    Scan inventory dir (excluding 'batch') for '<docid>_####.png' and return sorted unique docids.
    """
    ids = set()
    for p in visuals_dir.glob("*.png"):
        if not p.is_file():
            continue
        m = DOCID_FILE_RE.match(p.name)
        if m:
            ids.add(int(m.group("docid")))
    return sorted(ids)


def validate_inventory_has_document(visuals_dir: Path, document_id: int) -> None:
    """
    Ensure at least one '<document_id>_####.png' exists in visuals_dir.
    Fail loudly if not, and print what docids are present.
    """
    matches = list(visuals_dir.glob(f"{document_id}_*.png"))
    ok = any(p.is_file() and DOCID_FILE_RE.match(p.name) for p in matches)
    if ok:
        return
    found = present_document_ids(visuals_dir)
    raise ValueError(
        f"No visuals found for --document_id={document_id} in --visuals_dir='{visuals_dir}'. "
        f"Found document_ids in this directory: {found if found else 'NONE'}. "
        f"Fix --document_id or point --visuals_dir to the correct manager inventory folder."
    )


def validate_ledger_matches_cfg(ledger: dict, cfg: 'RunConfig', ledger_path: Path) -> None:
    """
    If ledger exists, ensure it matches the CLI manager_id and document_id to prevent cross-contamination.
    """
    if not ledger:
        return
    led_mid = ledger.get("manager_id")
    if led_mid is not None and int(led_mid) != cfg.manager_id:
        raise ValueError(
            f"Existing ledger JSON '{ledger_path.name}' has manager_id={led_mid}, but CLI --manager_id={cfg.manager_id}."
        )

    docs = ledger.get("documents") or []
    if docs:
        led_docid = docs[0].get("document_id")
        if led_docid is not None and int(led_docid) != cfg.document_id:
            raise ValueError(
                f"Existing ledger JSON '{ledger_path.name}' has document_id={led_docid}, but CLI --document_id={cfg.document_id}."
            )


@dataclass
class RunConfig:
    input_path: Path
    manager: str
    manager_id: int
    document_id: int
    document_date: str
    document_source: str

    visuals_dir: Path
    batch_size: int = 15

    model: str = "gpt-5.2"
    temperature: float = 0.1
    prompt_version: str = "rawvisual_v1"

    output_root: Optional[Path] = None  # computed


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def choose_output_root(input_path: Path) -> Path:
    """
    Use <HOME>/Downloads if it exists; else input file directory.
    """
    home = Path.home()
    downloads = home / "Downloads"
    if downloads.exists() and downloads.is_dir():
        return downloads
    return input_path.parent


def clamp_batch(n: int) -> int:
    if n <= 0:
        return MAX_BATCH
    return min(n, MAX_BATCH)


def data_url_for_png(path: Path) -> str:
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{b64}"


# -----------------------------
# JSON ledger (Option A)
# -----------------------------
def load_ledger(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def processed_filenames(ledger: Dict[str, Any]) -> Set[str]:
    """
    Return set of filenames already interpreted.
    We treat items in ledger['visuals'] as done if they have a non-empty interpretation.
    """
    done: Set[str] = set()
    visuals = ledger.get("visuals", [])
    for v in visuals:
        fn = v.get("filename")
        interp = v.get("interpretation")
        if fn and interp:
            done.add(fn)
    return done


def init_ledger_if_needed(cfg: RunConfig, ledger_path: Path) -> Dict[str, Any]:
    led = load_ledger(ledger_path)
    if led:
        return led

    extraction_dt = now_iso()
    return {
        "schema_version": cfg.prompt_version,
        "manager_id": cfg.manager_id,
        "manager_name": cfg.manager,
        "documents": [
            {
                "document_id": cfg.document_id,
                "document_date": cfg.document_date,
                "document_source": cfg.document_source,
                "extraction_datetime": extraction_dt,
            }
        ],
        "visual_directory": str(cfg.visuals_dir.resolve()),
        "visuals": [],  # append-only list
        "coverage": {
            "visual_count": 0,
            "interpreted_count": 0,
            "run_datetime": extraction_dt,
        },
    }


def update_coverage(ledger: Dict[str, Any]) -> None:
    visuals = ledger.get("visuals", [])
    ledger["coverage"]["visual_count"] = len(visuals)
    ledger["coverage"]["interpreted_count"] = sum(1 for v in visuals if v.get("interpretation"))
    ledger["coverage"]["run_datetime"] = now_iso()


# -----------------------------
# Batch selection + staging
# -----------------------------
def list_inventory_images(cfg: RunConfig) -> List[Path]:
    """
    Inventory = cfg.visuals_dir (excluding batch folder).
    Only include files matching <document_id>_####.png.
    """
    inv = []
    for p in cfg.visuals_dir.glob(f"{cfg.document_id}_*.png"):
        if p.is_file() and DOCID_FILE_RE.match(p.name):
            inv.append(p)
    return sorted(inv, key=lambda x: x.name.lower())


def ensure_batch_dir(visuals_dir: Path) -> Path:
    b = visuals_dir / "batch"
    b.mkdir(parents=True, exist_ok=True)
    return b


def clear_batch_dir(batch_dir: Path) -> None:
    for p in batch_dir.glob("*"):
        try:
            if p.is_file():
                p.unlink()
            elif p.is_dir():
                shutil.rmtree(p)
        except Exception:
            pass


def stage_next_batch(cfg: RunConfig, ledger: Dict[str, Any]) -> Tuple[List[Path], Path]:
    """
    Copies the next N unprocessed inventory images into batch/ and returns their paths in batch.
    """
    done = processed_filenames(ledger)
    inventory = list_inventory_images(cfg)
    todo = [p for p in inventory if p.name not in done]

    selected = todo[:clamp_batch(cfg.batch_size)]

    batch_dir = ensure_batch_dir(cfg.visuals_dir)
    clear_batch_dir(batch_dir)

    staged: List[Path] = []
    for p in selected:
        dest = batch_dir / p.name
        shutil.copy2(p, dest)
        staged.append(dest)

    return staged, batch_dir


# -----------------------------
# GPT call (batch)
# -----------------------------
def build_prompt() -> str:
    return (
        "You are an investment analyst. Interpret extracted tables/charts from an investment deck. "
        "For each image, produce IC-ready notes: what it is, what it shows, key insights with IC takeaways. "
        "Do NOT invent numbers you cannot read. If uncertain, say 'unclear'. "
        "Return STRICT JSON in the format:\n"
        "{\n"
        '  "items": [\n'
        "    {\n"
        '      "filename": "55666_0001.png",\n'
        '      "title": "",\n'
        '      "what_it_is": "",\n'
        '      "components": ["..."],\n'
        '      "insights": [{"headline":"","detail":"","ic_takeaway":""}],\n'
        '      "one_sentence_takeaway": "",\n'
        '      "reading_notes": ""\n'
        "    }\n"
        "  ]\n"
        "}\n"
    )



def _extract_first_json_object(text: str) -> dict:
    """
    Best-effort: parse JSON from model output even if it includes extra text.
    """
    text = text.strip()
    try:
        return json.loads(text)
    except Exception:
        pass

    start = text.find("{")
    if start == -1:
        raise ValueError("No JSON object found in model output.")

    stack = 0
    for i in range(start, len(text)):
        ch = text[i]
        if ch == "{":
            stack += 1
        elif ch == "}":
            stack -= 1
            if stack == 0:
                candidate = text[start:i+1]
                return json.loads(candidate)

    raise ValueError("Could not extract a complete JSON object from model output.")

def interpret_batch(
    client: OpenAI,
    cfg: RunConfig,
    batch_images: List[Path],
) -> Dict[str, Any]:
    """
    Sends up to 15 images in one call and returns parsed JSON.

    Note: Some OpenAI Python SDK versions don't support `response_format` on `responses.create`.
    We therefore request STRICT JSON in the prompt and then extract/parse it robustly.
    """
    content: List[Dict[str, Any]] = [{"type": "input_text", "text": build_prompt()}]
    for img in batch_images:
        content.append({"type": "input_image", "image_url": data_url_for_png(img)})

    resp = client.responses.create(
        model=cfg.model,
        input=[{"role": "user", "content": content}],
        temperature=cfg.temperature,
    )

    return _extract_first_json_object(resp.output_text)

# -----------------------------
# DOCX append
# -----------------------------
def load_or_create_docx(path: Path) -> Document:
    if path.exists():
        return Document(str(path))
    doc = Document()
    doc.add_heading("RAW Visual Interpretations", level=1)
    return doc


def append_to_docx(doc: Document, cfg: RunConfig, item: Dict[str, Any], rel_path: str) -> None:
    fn = item.get("filename", "")
    doc.add_page_break()
    doc.add_heading(f"{fn}", level=2)

    doc.add_paragraph(f"Manager: {cfg.manager} (ID: {cfg.manager_id})")
    doc.add_paragraph(f"Document: {cfg.document_id} | {cfg.document_date} | {cfg.document_source}")
    doc.add_paragraph(f"Visual file: {rel_path}")

    title = item.get("title") or ""
    if title:
        doc.add_paragraph(f"Title: {title}")

    doc.add_heading("What this is", level=3)
    doc.add_paragraph(item.get("what_it_is", ""))

    comps = item.get("components") or []
    if comps:
        doc.add_heading("Components", level=3)
        for c in comps:
            doc.add_paragraph(str(c), style="List Bullet")

    ins = item.get("insights") or []
    if ins:
        doc.add_heading("Key insights", level=3)
        for i in ins:
            doc.add_paragraph(i.get("headline", ""), style="List Bullet")
            if i.get("detail"):
                doc.add_paragraph(f"Detail: {i['detail']}")
            if i.get("ic_takeaway"):
                doc.add_paragraph(f"IC takeaway: {i['ic_takeaway']}")

    doc.add_heading("One-sentence takeaway", level=3)
    doc.add_paragraph(item.get("one_sentence_takeaway", ""))

    notes = item.get("reading_notes", "")
    if notes:
        doc.add_heading("Reading notes / caveats", level=3)
        doc.add_paragraph(notes)


# -----------------------------
# Main
# -----------------------------
def main() -> None:
    ap = argparse.ArgumentParser(description="Batch interpret extracted visuals and append to rawvisual JSON/DOCX (max 15 per run).")
    ap.add_argument("--input", type=Path, required=True, help="Original source document path (used for metadata + output-root selection).")
    ap.add_argument("--manager", type=str, required=True)
    ap.add_argument("--manager_id", type=int, required=True)
    ap.add_argument("--document_id", type=int, required=True)
    ap.add_argument("--document_date", type=str, required=True, help="YYYY-MM-DD")
    ap.add_argument("--document_source", type=str, required=True, help="Manager | Townsend | 3rd party")

    ap.add_argument("--visuals_dir", type=Path, required=True, help="Manager visual inventory dir (contains all visuals).")
    ap.add_argument("--batch_size", type=int, default=15, help="How many visuals to interpret this run (capped at 15). Default 15.")
    ap.add_argument("--model", type=str, default="gpt-5.2")
    ap.add_argument("--temperature", type=float, default=0.1)

    args = ap.parse_args()

    cfg = RunConfig(
        input_path=args.input,
        manager=args.manager,
        manager_id=args.manager_id,
        document_id=args.document_id,
        document_date=args.document_date,
        document_source=args.document_source,
        visuals_dir=args.visuals_dir,
        batch_size=clamp_batch(args.batch_size),
        model=args.model,
        temperature=args.temperature,
    )

    if not cfg.input_path.exists():
        raise FileNotFoundError(f"--input not found: {cfg.input_path}")
    if not cfg.visuals_dir.exists():
        raise FileNotFoundError(f"--visuals_dir not found: {cfg.visuals_dir}")

    # Safety validations to avoid CLI/inventory mismatches
    validate_visuals_dir_manager_id(cfg.visuals_dir, cfg.manager_id)
    validate_inventory_has_document(cfg.visuals_dir, cfg.document_id)

    cfg.output_root = choose_output_root(cfg.input_path)

    out_json = cfg.output_root / f"{cfg.manager_id}_rawvisual_{cfg.document_id}.json"
    out_docx = cfg.output_root / f"{cfg.manager_id}_rawvisual_{cfg.document_id}.docx"

    ledger = init_ledger_if_needed(cfg, out_json)
    validate_ledger_matches_cfg(ledger, cfg, out_json)

    staged, batch_dir = stage_next_batch(cfg, ledger)
    if not staged:
        update_coverage(ledger)
        out_json.write_text(json.dumps(ledger, indent=2), encoding="utf-8")
        print("No new visuals to interpret for this document.")
        print(f"JSON:  {out_json.resolve()}")
        print(f"DOCX:  {out_docx.resolve()}")
        return

    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set. In CMD: set OPENAI_API_KEY=YOUR_KEY")

    client = OpenAI(api_key=api_key)

    result = interpret_batch(client, cfg, staged)
    items = result.get("items", [])
    if not isinstance(items, list):
        raise RuntimeError("Model returned invalid JSON: missing items[] list.")

    by_fn: Dict[str, Dict[str, Any]] = {}
    for it in items:
        fn = it.get("filename")
        if fn:
            by_fn[fn] = it

    doc = load_or_create_docx(out_docx)
    run_dt = now_iso()

    appended = 0
    for img in staged:
        fn = img.name
        it = by_fn.get(fn)
        if not it:
            it = {
                "filename": fn,
                "title": "",
                "what_it_is": "",
                "components": [],
                "insights": [],
                "one_sentence_takeaway": "",
                "reading_notes": "No interpretation returned for this image in the batch response.",
            }

        inv_path = cfg.visuals_dir / fn
        rel_path = str(inv_path.resolve())

        record = {
            "visual_id": f"V{len(ledger.get('visuals', [])) + 1:04d}",
            "filename": fn,
            "relative_path": rel_path,
            "visual_refs": [
                {
                    "document_id": cfg.document_id,
                    "filename": fn,
                    "relative_path": rel_path,
                }
            ],
            "provenance": {
                "document_id": cfg.document_id,
                "document_date": cfg.document_date,
                "document_source": cfg.document_source,
                "extraction_datetime": ledger["documents"][0]["extraction_datetime"],
            },
            "assigned_by": {
                "model": cfg.model,
                "prompt_version": cfg.prompt_version,
                "run_datetime": run_dt,
            },
            "interpretation": {
                "title": it.get("title", ""),
                "what_it_is": it.get("what_it_is", ""),
                "components": it.get("components", []),
                "insights": it.get("insights", []),
                "one_sentence_takeaway": it.get("one_sentence_takeaway", ""),
                "reading_notes": it.get("reading_notes", ""),
            },
        }

        ledger["visuals"].append(record)
        append_to_docx(doc, cfg, it, rel_path)
        appended += 1

    update_coverage(ledger)
    out_json.write_text(json.dumps(ledger, indent=2), encoding="utf-8")
    doc.save(str(out_docx))

    clear_batch_dir(batch_dir)

    print(f"Interpreted {appended} visuals (batch_size={cfg.batch_size}, cap={MAX_BATCH}).")
    print(f"JSON:  {out_json.resolve()}")
    print(f"DOCX:  {out_docx.resolve()}")


if __name__ == "__main__":
    main()
