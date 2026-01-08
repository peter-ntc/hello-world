#!/usr/bin/env python3
"""
MCAF Aggregator v1

Purpose:
- Merge a single per-document MCAF JSON (e.g., 10067_mcaf_55666.json) into an existing
  per-manager aggregate store (e.g., 10067_aggregate.json/.docx).

Key properties:
- Conservative, lossless accumulation (no synthesis, no conflict resolution)
- Provenance preserved per bullet occurrence
- Idempotent (re-running the same input does not duplicate bullets)
- DOCX is a rendered view from aggregate JSON; controlled via DOCX_WRITE (default Yes)

Default output directory:
- C:/Users/prash/Downloads/ if it exists; otherwise uses the directory of aggregate_json.

Note: We use forward slashes in this docstring to avoid Python backslash escape issues.
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Tuple, Any

from docx import Document


# ---------- MCAF SECTION DEFINITIONS (top-level only) ----------

MCAF_SECTIONS: List[Tuple[str, str]] = [
    ("S01", "Team Experience & Structure"),
    ("S02", "Organizational Depth & Infrastructure"),
    ("S03", "Sector-Specific Capabilities"),
    ("S04", "Investment Process & Governance"),
    ("S05", "Alignment of Interest"),
    ("S06", "Fiduciary Standards & Reputation"),
    ("S07", "Track Record Quality"),
    ("S08", "Platform Evolution & Stability"),
    ("S09", "Vertical Integration & Operating Edge"),
    ("S10", "Conflicts & Transparency"),
    ("S11", "Institutional Readiness"),
    ("S12", "Fee & Incentive Assessment"),
    ("S13", "Team Culture & Retention"),
    ("S14", "Overall Manager Rating"),
    ("S15", "Other Comprehensive Manager Information"),
]

SECTION_TITLES = dict(MCAF_SECTIONS)

DEFAULT_OUTPUT_DIR = r"C:\Users\prash\Downloads"

SOURCE_PRIORITY = {
    "Townsend": 0,
    "Manager": 1,
    "3rd party": 2,
}

def parse_yes_no(val: str) -> bool:
    v = (val or "").strip().lower()
    if v in ("y", "yes", "true", "1", "on"):
        return True
    if v in ("n", "no", "false", "0", "off"):
        return False
    raise argparse.ArgumentTypeError("DOCX_WRITE must be Yes/No (or True/False).")

def resolve_output_dir(arg_output_dir: str | None, aggregate_json_path: Path) -> Path:
    if arg_output_dir:
        out = Path(arg_output_dir).expanduser().resolve()
        out.mkdir(parents=True, exist_ok=True)
        return out

    default = Path(DEFAULT_OUTPUT_DIR)
    if default.exists() and default.is_dir():
        return default.resolve()

    return aggregate_json_path.parent.resolve()

def now_utc_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()

def normalize_text(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = s.rstrip(" .;:,")
    return s

def ensure_source_canonical(src: str) -> str:
    s = (src or "").strip().lower()
    if s in ("t", "townsend"):
        return "Townsend"
    if s in ("m", "manager"):
        return "Manager"
    if s in ("3", "3rd", "3rd party", "third party", "third-party"):
        return "3rd party"
    if "3" in s and "party" in s:
        return "3rd party"
    return (src or "").strip() or "3rd party"

def safe_parse_date(d: str) -> datetime:
    # expects YYYY-MM-DD; fall back to epoch if invalid
    try:
        return datetime.strptime(d, "%Y-%m-%d")
    except Exception:
        return datetime(1900, 1, 1)


@dataclass(frozen=True)
class BulletKey:
    document_id: int
    section_id: str
    raw_id: str

    def uid(self) -> str:
        return f"{self.document_id}:{self.section_id}:{self.raw_id}"


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def validate_aggregate_schema(agg: dict) -> None:
    if "manager_id" not in agg:
        raise ValueError("Aggregate JSON missing manager_id")
    if "sections" not in agg or not isinstance(agg["sections"], dict):
        raise ValueError("Aggregate JSON missing sections object")
    for sid, _ in MCAF_SECTIONS:
        if sid not in agg["sections"]:
            raise ValueError(f"Aggregate JSON missing section {sid}")
        if "bullets" not in agg["sections"][sid]:
            raise ValueError(f"Aggregate JSON section {sid} missing bullets array")


def validate_mcaf_schema(mcaf: dict) -> None:
    if "manager_id" not in mcaf:
        raise ValueError("MCAF JSON missing manager_id")
    if "sections" not in mcaf or not isinstance(mcaf["sections"], dict):
        raise ValueError("MCAF JSON missing sections object")
    for sid, _ in MCAF_SECTIONS:
        if sid not in mcaf["sections"]:
            raise ValueError(f"MCAF JSON missing section {sid}")
        if "bullets" not in mcaf["sections"][sid]:
            raise ValueError(f"MCAF JSON section {sid} missing bullets array")


def add_document_if_missing(agg: dict, doc_meta: dict) -> bool:
    """Returns True if added, False if already present."""
    doc_id = doc_meta.get("document_id")
    if doc_id is None:
        return False
    docs = agg.setdefault("documents", [])
    for d in docs:
        if d.get("document_id") == doc_id:
            return False
    docs.append(doc_meta)
    return True


def build_doc_meta_from_mcaf(mcaf: dict) -> dict:
    docs = mcaf.get("documents") or []
    if docs:
        # per-doc filer writes one document entry
        d0 = docs[0] or {}
        return {
            "document_id": d0.get("document_id"),
            "document_date": d0.get("document_date"),
            "document_source": ensure_source_canonical(d0.get("document_source", "")),
            "extraction_datetime": d0.get("extraction_datetime"),
        }

    # fallback: infer from first bullet provenance if needed
    for sid, _ in MCAF_SECTIONS:
        bullets = (mcaf.get("sections", {}).get(sid, {}) or {}).get("bullets", []) or []
        if bullets:
            prov = (bullets[0] or {}).get("provenance", {}) or {}
            return {
                "document_id": prov.get("document_id"),
                "document_date": prov.get("document_date"),
                "document_source": ensure_source_canonical(prov.get("document_source", "")),
                "extraction_datetime": prov.get("extraction_datetime"),
            }
    return {"document_id": None, "document_date": None, "document_source": None, "extraction_datetime": None}


def index_existing_bullets(agg: dict) -> Dict[str, set]:
    """
    Build per-section sets of existing bullet_uids for idempotency.
    Returns: {section_id: set(bullet_uid)}
    """
    idx: Dict[str, set] = {sid: set() for sid, _ in MCAF_SECTIONS}
    for sid, _ in MCAF_SECTIONS:
        for b in agg["sections"][sid].get("bullets", []) or []:
            uid = b.get("bullet_uid")
            if uid:
                idx[sid].add(uid)
            else:
                # support older aggregates if needed: derive uid if fields exist
                prov = b.get("provenance", {}) or {}
                doc_id = prov.get("document_id")
                raw_id = b.get("raw_id")
                if doc_id is not None and raw_id:
                    idx[sid].add(f"{doc_id}:{sid}:{raw_id}")
    return idx


def merge_mcaf_into_aggregate(agg: dict, mcaf: dict, mcaf_path: Path) -> dict:
    """
    Merges mcaf sections into agg sections. Updates coverage and returns a run audit dict.
    """
    run_dt = now_utc_iso()
    audit_run = {
        "run_datetime": run_dt,
        "mcaf_json": str(mcaf_path),
        "document": {},
        "added": {"documents_added": 0, "bullets_added_by_section": {sid: 0 for sid, _ in MCAF_SECTIONS}},
        "skipped": {"duplicate_bullets_by_section": {sid: 0 for sid, _ in MCAF_SECTIONS}},
        "totals_after": {},
    }

    # Enforce manager match
    agg_mid = int(agg.get("manager_id"))
    mcaf_mid = int(mcaf.get("manager_id"))
    if agg_mid != mcaf_mid:
        raise ValueError(f"Manager ID mismatch: aggregate={agg_mid} vs mcaf={mcaf_mid}")

    # Populate manager_name if missing
    if not agg.get("manager_name") and mcaf.get("manager_name"):
        agg["manager_name"] = mcaf.get("manager_name")

    # Document meta
    doc_meta = build_doc_meta_from_mcaf(mcaf)
    audit_run["document"] = doc_meta

    if add_document_if_missing(agg, doc_meta):
        audit_run["added"]["documents_added"] = 1

    existing_uid_by_section = index_existing_bullets(agg)

    # Merge bullets per section
    for sid, _ in MCAF_SECTIONS:
        incoming = mcaf["sections"][sid].get("bullets", []) or []
        target_list = agg["sections"][sid].setdefault("bullets", [])

        # document-aware safety dedupe: within same doc+section, dedupe by normalized text
        # Build set from existing bullets in this section
        existing_norm_by_doc = set()
        for b in target_list:
            prov = b.get("provenance", {}) or {}
            doc_id = prov.get("document_id")
            txt_norm = b.get("text_norm") or normalize_text(b.get("text", ""))
            if doc_id is not None and txt_norm:
                existing_norm_by_doc.add((doc_id, txt_norm))

        for b in incoming:
            # Required fields
            raw_id = b.get("raw_id")
            text = b.get("text", "")
            prov = b.get("provenance", {}) or {}
            doc_id = prov.get("document_id")
            if doc_id is None or not raw_id:
                # skip malformed bullets but record as skipped (could be elevated later)
                audit_run["skipped"]["duplicate_bullets_by_section"][sid] += 1
                continue

            key = BulletKey(document_id=int(doc_id), section_id=sid, raw_id=str(raw_id))
            uid = key.uid()

            # Idempotency
            if uid in existing_uid_by_section[sid]:
                audit_run["skipped"]["duplicate_bullets_by_section"][sid] += 1
                continue

            # Safety dedupe within same document+section by normalized text
            txt_norm = normalize_text(text)
            if (int(doc_id), txt_norm) in existing_norm_by_doc:
                # same doc+section, likely repeat; skip
                audit_run["skipped"]["duplicate_bullets_by_section"][sid] += 1
                continue

            # Build aggregate bullet entry (preserve payload; add bullet_uid + text_norm for future)
            out_b = dict(b)  # shallow copy
            out_b["bullet_uid"] = uid
            out_b["text_norm"] = txt_norm
            # ensure canonical source in provenance
            out_b.setdefault("provenance", {})
            out_b["provenance"]["document_source"] = ensure_source_canonical(out_b["provenance"].get("document_source", ""))

            target_list.append(out_b)

            existing_uid_by_section[sid].add(uid)
            existing_norm_by_doc.add((int(doc_id), txt_norm))
            audit_run["added"]["bullets_added_by_section"][sid] += 1

    # Update coverage
    docs = agg.get("documents") or []
    agg.setdefault("coverage", {})
    agg["coverage"]["document_count"] = len(docs)

    # Count bullet occurrences + unique bullet_uids
    total_occ = 0
    uid_set = set()
    for sid, _ in MCAF_SECTIONS:
        for b in agg["sections"][sid].get("bullets", []) or []:
            total_occ += 1
            uid = b.get("bullet_uid")
            if uid:
                uid_set.add(uid)

    agg["coverage"]["bullet_occurrence_count"] = total_occ
    agg["coverage"]["unique_bullet_uid_count"] = len(uid_set)
    agg["last_updated_datetime"] = run_dt

    audit_run["totals_after"] = {
        "document_count": agg["coverage"]["document_count"],
        "bullet_occurrence_count": agg["coverage"]["bullet_occurrence_count"],
        "unique_bullet_uid_count": agg["coverage"]["unique_bullet_uid_count"],
    }

    return audit_run


def render_aggregate_docx(agg: dict, output_path: Path) -> None:
    doc = Document()

    manager_id = agg.get("manager_id", "")
    manager_name = agg.get("manager_name", "")
    doc.add_heading(f"MCAF Aggregate: {manager_name} (ID {manager_id})".strip(), level=0)

    doc.add_paragraph(f"Manager ID: {manager_id}")
    if manager_name:
        doc.add_paragraph(f"Manager Name: {manager_name}")
    doc.add_paragraph(f"Documents in aggregate: {agg.get('coverage', {}).get('document_count', 0)}")
    doc.add_paragraph(f"Last updated: {agg.get('last_updated_datetime', '')}")

    # Render in S01..S15 order
    for sid, title in MCAF_SECTIONS:
        doc.add_heading(f"{sid}: {title}", level=1)
        items = agg["sections"][sid].get("bullets", []) or []

        def sort_key(it: dict):
            prov = it.get("provenance", {}) or {}
            d = prov.get("document_date", "1900-01-01")
            dt = safe_parse_date(d)
            src = ensure_source_canonical(prov.get("document_source", ""))
            docid = prov.get("document_id") or 0
            raw_id = it.get("raw_id") or ""
            return (-int(dt.timestamp()), SOURCE_PRIORITY.get(src, 99), int(docid), str(raw_id))

        for it in sorted(items, key=sort_key):
            prov = it.get("provenance", {}) or {}
            d = prov.get("document_date", "")
            src = ensure_source_canonical(prov.get("document_source", ""))
            docid = prov.get("document_id", "")
            rid = it.get("raw_id", "")
            text = it.get("text", "")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
            p.add_run(f"  [{d} | {src} | Doc {docid} | {rid}]").italic = True

    doc.save(str(output_path))


def update_audit_file(audit_path: Path, audit_run: dict) -> None:
    if audit_path.exists():
        try:
            existing = load_json(audit_path)
        except Exception:
            existing = {}
    else:
        existing = {}

    history = existing.get("history")
    if not isinstance(history, list):
        history = []
    history.append(audit_run)

    existing["schema_version"] = "mcaf_aggregate_audit_v1"
    existing["history"] = history
    existing["last_run_datetime"] = audit_run.get("run_datetime")

    with audit_path.open("w", encoding="utf-8") as f:
        json.dump(existing, f, indent=2, ensure_ascii=False)


def main() -> None:
    parser = argparse.ArgumentParser(description="Merge a per-document MCAF JSON into a per-manager aggregate store.")
    parser.add_argument("--mcaf_json", required=True, help="Path to per-document MCAF JSON (e.g., 10067_mcaf_55666.json)")
    parser.add_argument("--aggregate_json", required=True, help="Path to existing aggregate JSON (e.g., 10067_aggregate.json)")
    parser.add_argument("--aggregate_docx", required=True, help="Path to existing aggregate DOCX (e.g., 10067_aggregate.docx)")
    parser.add_argument(
        "--output_dir",
        default=None,
        help=r"Output directory (default: C:\Users\prash\Downloads if it exists; else aggregate_json directory)",
    )
    parser.add_argument(
        "--DOCX_WRITE",
        type=parse_yes_no,
        default=True,
        help="Whether to write the updated DOCX (Yes/No). Default: Yes",
    )
    args = parser.parse_args()

    mcaf_path = Path(args.mcaf_json).expanduser().resolve()
    agg_json_path = Path(args.aggregate_json).expanduser().resolve()
    agg_docx_path = Path(args.aggregate_docx).expanduser().resolve()

    if not mcaf_path.exists():
        raise FileNotFoundError(f"MCAF JSON not found: {mcaf_path}")
    if not agg_json_path.exists():
        raise FileNotFoundError(f"Aggregate JSON not found: {agg_json_path}")
    if not agg_docx_path.exists() and args.DOCX_WRITE:
        raise FileNotFoundError(f"Aggregate DOCX not found: {agg_docx_path}")

    mcaf = load_json(mcaf_path)
    agg = load_json(agg_json_path)

    validate_mcaf_schema(mcaf)
    validate_aggregate_schema(agg)

    # Manager ID match and canonical output names
    mcaf_mid = int(mcaf.get("manager_id"))
    agg_mid = int(agg.get("manager_id"))
    if mcaf_mid != agg_mid:
        raise ValueError(f"Manager ID mismatch: aggregate={agg_mid} vs mcaf={mcaf_mid}")

    out_dir = resolve_output_dir(args.output_dir, agg_json_path)
    out_dir.mkdir(parents=True, exist_ok=True)

    base = f"{mcaf_mid}_aggregate"
    out_json = out_dir / f"{base}.json"
    out_docx = out_dir / f"{base}.docx"
    out_audit = out_dir / f"{base}_audit.json"

    audit_run = merge_mcaf_into_aggregate(agg, mcaf, mcaf_path)

    # Write updated aggregate JSON (canonical filename)
    with out_json.open("w", encoding="utf-8") as f:
        json.dump(agg, f, indent=2, ensure_ascii=False)

    # Update audit
    update_audit_file(out_audit, audit_run)

    # Write DOCX (optional) rendered from updated aggregate JSON
    if args.DOCX_WRITE:
        render_aggregate_docx(agg, out_docx)

    print("Wrote:")
    print(f"  {out_json}")
    print(f"  {out_audit}")
    if args.DOCX_WRITE:
        print(f"  {out_docx}")
    else:
        print("  (DOCX_WRITE=No; DOCX not written)")


if __name__ == "__main__":
    main()
